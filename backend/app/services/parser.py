import PyPDF2
import docx
import logging
from typing import Optional

logger = logging.getLogger(__name__)

class DocumentParser:
    """Enhanced document parser with better error handling and validation"""
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF with comprehensive error handling
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If PDF cannot be read or is corrupted
        """
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF is encrypted: {file_path}")
                    try:
                        # Try to decrypt with empty password
                        pdf_reader.decrypt('')
                    except:
                        raise ValueError("PDF is password-protected and cannot be read")
                
                # Check if PDF has pages
                num_pages = len(pdf_reader.pages)
                if num_pages == 0:
                    raise ValueError("PDF has no pages")
                
                logger.info(f"Processing PDF with {num_pages} pages")
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"No text extracted from page {page_num + 1}")
                            
                    except Exception as e:
                        logger.error(f"Error extracting page {page_num + 1}: {e}")
                        continue
                
                # Check if any text was extracted
                if not text.strip():
                    raise ValueError("PDF appears to be empty or scanned (no text could be extracted)")
                        
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error for {file_path}: {e}")
            raise ValueError(f"Corrupted or invalid PDF file: {str(e)}")
        except Exception as e:
            logger.error(f"Failed to read PDF {file_path}: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX with error handling
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If DOCX cannot be read
        """
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                raise ValueError("DOCX file appears to be empty")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to read DOCX {file_path}: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file with encoding handling
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text as string
        """
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read().strip()
                    if text:
                        return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT file with {encoding}: {e}")
                continue
        
        raise ValueError("Could not read text file with any supported encoding")
    
    @staticmethod
    def validate_file(file_path: str, file_extension: str) -> None:
        """
        Validate file before processing
        
        Args:
            file_path: Path to file
            file_extension: File extension
            
        Raises:
            ValueError: If file validation fails
        """
        import os
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise ValueError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError("File is empty")
        
        if file_size > DocumentParser.MAX_FILE_SIZE:
            raise ValueError(f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds maximum allowed size (10MB)")
        
        # Validate extension
        allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
        if file_extension.lower() not in allowed_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}. Allowed: {', '.join(allowed_extensions)}")
        
        logger.info(f"File validation passed: {file_path} ({file_size / 1024:.2f}KB)")
    
    @staticmethod
    def extract_text(file_path: str, file_extension: str) -> str:
        """
        Extract text from file with validation and error handling
        
        Args:
            file_path: Path to file
            file_extension: File extension (.pdf, .docx, .txt)
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file cannot be processed
        """
        # Validate file first
        DocumentParser.validate_file(file_path, file_extension)
        
        # Extract text based on file type
        try:
            if file_extension.lower() == '.pdf':
                text = DocumentParser.extract_text_from_pdf(file_path)
            elif file_extension.lower() in ['.docx', '.doc']:
                text = DocumentParser.extract_text_from_docx(file_path)
            elif file_extension.lower() == '.txt':
                text = DocumentParser.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Validate extracted text
            if not text or len(text.strip()) < 50:
                raise ValueError("Insufficient text extracted from document (less than 50 characters)")
            
            logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
            return text
            
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error processing {file_path}: {e}")
            raise ValueError(f"Failed to process document: {str(e)}")
        